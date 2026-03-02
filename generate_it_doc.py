"""
Generates the IT Admin Slack App Response Word document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def set_heading(paragraph, text, level=1):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return paragraph

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_heading(p, text, level)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_code_block(doc, text):
    """Add a monospace code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header background
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E74B5")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = val
            if r_idx % 2 == 1:
                tc = row_cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EBF3FB")
                tcPr.append(shd)
    return table

def add_checklist_item(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("☐  " + text)
    return p

# ---------------------------------------------------------------------------
# Document Title
# ---------------------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_run = title.add_run("Slack Application Security Review")
t_run.bold = True
t_run.font.size = Pt(20)
t_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
s_run = subtitle.add_run("IT Admin Response — ClaudeResearchApp")
s_run.font.size = Pt(12)
s_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()  # spacer

# ---------------------------------------------------------------------------
# Q1: Data Flow
# ---------------------------------------------------------------------------
add_heading(doc, "1.  Data Flow Diagram", level=2)

add_body(doc, (
    "ClaudeResearchApp is a local CLI tool (context.py) that searches Slack for internal "
    "messages mentioning a prospect company name. The flow is strictly read-only and "
    "one-directional — no data is written back to Slack at any point."
))

doc.add_paragraph()

flow = (
    "[Nick's Laptop — CLI]\n"
    "      │  python context.py \"Acme Corp\"\n"
    "      │  (reads SLACK_USER_TOKEN from .env)\n"
    "      ▼\n"
    "[slack_sdk WebClient]\n"
    "      │  HTTPS → Slack search.messages API\n"
    "      │  query: \"Acme Corp\", up to 20 results, newest first\n"
    "      ▼\n"
    "[Slack API — api.slack.com]\n"
    "      │  Returns: channel name, username, message text (250 chars),\n"
    "      │           timestamp, permalink\n"
    "      ▼\n"
    "[context.py — in memory only]\n"
    "      │  Combined with Salesforce, Gmail, Drive, and Notion data\n"
    "      ▼\n"
    "[Anthropic Claude API — api.anthropic.com]\n"
    "      │  HTTPS — full context prompt sent, synthesized briefing returned\n"
    "      ▼\n"
    "[Local Disk — ~/Prospecting/reports/]\n"
    "      └─ {company}_{date}_internal_context.md"
)
add_code_block(doc, flow)

doc.add_paragraph()
add_body(doc, (
    "Slack access scope: search:read on a User OAuth Token (xoxp-...) — read-only, "
    "searches only channels Nick's own Slack account can already see. "
    "No admin access, no message deletion, no posting capability."
))

doc.add_paragraph()

# ---------------------------------------------------------------------------
# Q2: Data Storage
# ---------------------------------------------------------------------------
add_heading(doc, "2.  Where Data Gets Stored at Each Step", level=2)

add_table(doc,
    headers=["Step", "Location", "What's Stored", "Persisted?"],
    rows=[
        ["Token",           ".env on Nick's laptop",          "SLACK_USER_TOKEN (xoxp-...)",                              "Yes — plaintext file"],
        ["In-flight",       "HTTPS request (TLS)",            "Company name search query",                                "No — transit only"],
        ["In memory",       "Python process",                 "Message text, channel, username, timestamp, permalink",    "No — cleared on exit"],
        ["Sent to Claude",  "Anthropic API servers",          "Message content + all other connector data",               "Per Anthropic API policy — not used for training by default"],
        ["Final report",    "~/Prospecting/reports/*.md",     "Claude's synthesized briefing (includes Slack summary)",   "Yes — plaintext markdown until manually deleted"],
    ]
)

doc.add_paragraph()
add_body(doc, (
    "No Slack data is stored in a database, cloud bucket, or shared server. "
    "Everything lives on Nick's local machine or passes transiently through the Anthropic API. "
    "Anthropic's API data usage policy is available at: https://www.anthropic.com/legal/usage-policy"
))

doc.add_paragraph()

# ---------------------------------------------------------------------------
# Q3: Responsible Party
# ---------------------------------------------------------------------------
add_heading(doc, "3.  Who Is Responsible for Managing the Application", level=2)

add_bullet(doc, "Nick — created and owns the app at api.slack.com under his Slack account", bold_prefix="App owner / technical contact: ")
add_bullet(doc, "ClaudeResearchApp (visible at api.slack.com/apps)", bold_prefix="App name: ")
add_bullet(doc, "User OAuth Token only — not a bot, not added to any channel, not visible to other users in Slack", bold_prefix="Installation type: ")
add_bullet(doc, "The xoxp- token is tied to Nick's Slack identity; it can only search channels Nick can already access", bold_prefix="Token is personal: ")
add_bullet(doc, "The app appears under Slack Admin → Manage Apps and can be restricted or revoked at any time without Nick's involvement", bold_prefix="IT Admin control: ")

doc.add_paragraph()

# ---------------------------------------------------------------------------
# Q4: Termination
# ---------------------------------------------------------------------------
add_heading(doc, "4.  How to Terminate the Application", level=2)

add_body(doc, "There are three options, from least to most complete:")
doc.add_paragraph()

add_heading(doc, "Option A — Revoke the token only (soft stop, app record stays)", level=3)
steps_a = [
    "Go to api.slack.com/apps → select ClaudeResearchApp",
    "Navigate to OAuth & Permissions",
    "Click Revoke Token next to the User OAuth Token",
    "Token is immediately invalidated — any further API call returns token_revoked",
    "Delete or blank out SLACK_USER_TOKEN in ~/Prospecting/.env",
]
for s in steps_a:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s)

doc.add_paragraph()
add_heading(doc, "Option B — Delete the app entirely (hard stop)", level=3)
steps_b = [
    "Go to api.slack.com/apps → select ClaudeResearchApp",
    "Scroll to the bottom of Basic Information",
    "Click Delete App — all tokens are immediately revoked and the app is removed from Slack entirely",
]
for s in steps_b:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s)

doc.add_paragraph()
add_heading(doc, "Option C — IT Admin revokes directly (no action from Nick required)", level=3)
steps_c = [
    "Slack Admin Console → Manage Apps",
    "Find ClaudeResearchApp → click Restrict or Remove",
    "Revokes the installation org-wide immediately",
]
for s in steps_c:
    p = doc.add_paragraph(style="List Number")
    p.add_run(s)

doc.add_paragraph()
add_heading(doc, "Post-Termination Cleanup Checklist", level=3)
add_checklist_item(doc, "Delete or blank SLACK_USER_TOKEN in ~/Prospecting/.env")
add_checklist_item(doc, "Delete any reports in ~/Prospecting/reports/ that contain sensitive content")
add_checklist_item(doc, "Optionally uninstall the Slack SDK: pip uninstall slack-sdk")

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
out_path = "/Users/nick/Prospecting/reports/ClaudeResearchApp_IT_Admin_Response.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
